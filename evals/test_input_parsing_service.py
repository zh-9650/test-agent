from __future__ import annotations

from pathlib import Path

from docx import Document
from PIL import Image

from core.design_studio.contracts import (
    Authority,
    ParseFidelityStatus,
    SourceInput,
)
from core.design_studio.parsing import InputParsingService
from core.design_studio.parsing import ParseFidelityGate


def _write_docx(path: Path, *, body: str = "正文") -> None:
    image_path = path.with_suffix(".png")
    Image.new("RGB", (2, 2), color="white").save(image_path)

    document = Document()
    document.add_heading("绩效规则", level=1)
    document.add_paragraph(body)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "角色"
    table.cell(0, 1).text = "动作"
    table.cell(1, 0).text = "员工"
    table.cell(1, 1).text = "提交"
    document.add_picture(str(image_path))
    document.save(path)


def test_docx_parse_preserves_inventory_locators_and_source_version(tmp_path: Path) -> None:
    path = tmp_path / "requirement.docx"
    _write_docx(path)
    service = InputParsingService.default()
    source = SourceInput(
        source_id="SRC-PRD-001",
        path=path,
        source_kind="requirement",
        authority=Authority.NORMATIVE,
        required=True,
    )

    first = service.parse(source)
    repeated = service.parse(source)

    assert first.fidelity_report.status == ParseFidelityStatus.COMPLETE
    assert first.fidelity_report.detected_inventory["tables"] == 1
    assert first.fidelity_report.detected_inventory["rows"] == 2
    assert first.fidelity_report.detected_inventory["cells"] == 4
    assert first.fidelity_report.detected_inventory["drawings"] == 1
    assert first.fidelity_report.detected_inventory["media"] == 1
    assert first.fidelity_report.detected_inventory == (
        first.fidelity_report.parsed_inventory
    )
    assert all(block.locator for block in first.blocks)
    assert all(block.source_hash == first.source.sha256 for block in first.blocks)
    assert [block.block_id for block in first.blocks] == [
        block.block_id for block in repeated.blocks
    ]

    _write_docx(path, body="正文已变更")
    changed = service.parse(source)

    assert changed.source.sha256 != first.source.sha256
    assert [block.block_id for block in changed.blocks] != [
        block.block_id for block in first.blocks
    ]


def test_markdown_parse_preserves_structures_without_flattening(tmp_path: Path) -> None:
    path = tmp_path / "api.md"
    path.write_text(
        "\n".join(
            [
                "# 绩效接口",
                "",
                "接口说明见[规范](https://example.test/spec)。",
                "",
                "| 字段 | 类型 |",
                "| --- | --- |",
                "| status | string |",
                "",
                "```json",
                '{"status": "draft"}',
                "```",
                "",
                "## 状态枚举",
                "",
                "状态必须可追溯。",
            ]
        ),
        encoding="utf-8",
    )

    result = InputParsingService.default().parse(
        SourceInput(
            source_id="SRC-API-001",
            path=path,
            source_kind="api_document",
            authority=Authority.TECHNICAL,
        )
    )

    assert result.fidelity_report.status == ParseFidelityStatus.COMPLETE
    assert result.fidelity_report.detected_inventory == {
        "headings": 2,
        "paragraphs": 2,
        "tables": 1,
        "table_rows": 2,
        "code_blocks": 1,
        "links": 1,
    }
    assert result.fidelity_report.detected_inventory == (
        result.fidelity_report.parsed_inventory
    )
    assert any(
        block.block_type == "code_block"
        and block.text_content == '{"status": "draft"}'
        and block.structured_content["language"] == "json"
        for block in result.blocks
    )
    assert any(
        block.block_type == "link"
        and block.structured_content["target"] == "https://example.test/spec"
        for block in result.blocks
    )


def test_openapi_yaml_parse_preserves_operations_schemas_enums_and_refs(
    tmp_path: Path,
) -> None:
    path = tmp_path / "openapi.yaml"
    path.write_text(
        """
openapi: 3.0.3
paths:
  /reviews:
    post:
      operationId: submitReview
      parameters:
        - name: employeeId
          in: query
          schema:
            type: string
      requestBody:
        content:
          application/json:
            schema:
              $ref: '#/components/schemas/Review'
      responses:
        '200':
          description: success
          content:
            application/json:
              schema:
                $ref: '#/components/schemas/Review'
components:
  schemas:
    Review:
      type: object
      properties:
        status:
          $ref: '#/components/schemas/Status'
    Status:
      type: string
      enum: [draft, submitted]
""".strip(),
        encoding="utf-8",
    )

    result = InputParsingService.default().parse(
        SourceInput(
            source_id="SRC-OPENAPI-001",
            path=path,
            source_kind="api_contract",
            authority=Authority.TECHNICAL,
        )
    )

    assert result.fidelity_report.status == ParseFidelityStatus.COMPLETE
    assert result.fidelity_report.detected_inventory == {
        "operations": 1,
        "parameters": 1,
        "request_bodies": 1,
        "responses": 1,
        "schemas": 2,
        "enums": 1,
        "refs": 3,
    }
    assert result.fidelity_report.detected_inventory == (
        result.fidelity_report.parsed_inventory
    )
    assert any(
        block.block_type == "api_operation"
        and block.structured_content["operation_id"] == "submitReview"
        for block in result.blocks
    )
    assert any(
        block.block_type == "schema"
        and block.locator == "#/components/schemas/Status"
        for block in result.blocks
    )


def test_openapi_unresolved_local_ref_is_failed_not_silently_ignored(
    tmp_path: Path,
) -> None:
    path = tmp_path / "broken.json"
    path.write_text(
        """
{
  "openapi": "3.0.3",
  "paths": {
    "/reviews": {
      "get": {
        "responses": {
          "200": {
            "description": "ok",
            "content": {
              "application/json": {
                "schema": {"$ref": "#/components/schemas/Missing"}
              }
            }
          }
        }
      }
    }
  },
  "components": {"schemas": {}}
}
""".strip(),
        encoding="utf-8",
    )

    result = InputParsingService.default().parse(
        SourceInput(
            source_id="SRC-OPENAPI-BROKEN",
            path=path,
            source_kind="api_contract",
            authority=Authority.TECHNICAL,
        )
    )

    assert result.fidelity_report.status == ParseFidelityStatus.FAILED
    assert "input.unresolved_ref" in {
        item.code for item in result.fidelity_report.errors
    }


def test_html_bundle_preserves_resource_closure_but_reports_dynamic_semantics_partial(
    tmp_path: Path,
) -> None:
    root = tmp_path / "prototype"
    root.mkdir()
    (root / "index.html").write_text(
        """
<!doctype html>
<html>
  <head>
    <link rel="stylesheet" href="styles.css">
  </head>
  <body>
    <h1>绩效管理</h1>
    <form>
      <input name="score">
      <button>保存</button>
    </form>
    <a href="detail.html">查看详情</a>
    <script src="app.js"></script>
  </body>
</html>
""".strip(),
        encoding="utf-8",
    )
    (root / "detail.html").write_text(
        '<main><span role="button" onclick="confirmReview()">确认</span></main>',
        encoding="utf-8",
    )
    (root / "styles.css").write_text(
        "body { color: black; background: url('icon.png'); }",
        encoding="utf-8",
    )
    (root / "icon.png").write_bytes(b"fixture-resource")
    (root / "app.js").write_text(
        "function confirmReview() { return true; }",
        encoding="utf-8",
    )

    result = InputParsingService.default().parse(
        SourceInput(
            source_id="SRC-PROTOTYPE-001",
            path=root,
            source_kind="prototype",
            authority=Authority.NORMATIVE,
            entry_points=["index.html", "detail.html"],
        )
    )

    assert result.fidelity_report.status == ParseFidelityStatus.PARTIAL
    assert result.fidelity_report.detected_inventory["entry_pages"] == 2
    assert result.fidelity_report.detected_inventory["direct_referenced_files"] == 4
    assert result.fidelity_report.detected_inventory["referenced_files"] == 5
    assert result.fidelity_report.detected_inventory["missing_references"] == 0
    assert result.fidelity_report.detected_inventory["native_controls"] == 3
    assert result.fidelity_report.detected_inventory["custom_controls"] == 1
    assert result.fidelity_report.detected_inventory == (
        result.fidelity_report.parsed_inventory
    )
    assert "input.rendered_interaction_semantics" in {
        item.code for item in result.fidelity_report.unsupported_features
    }
    assert {Path(item.local_path).name for item in result.related_sources} == {
        "index.html",
        "detail.html",
        "styles.css",
        "app.js",
        "icon.png",
    }
    manifest = next(
        block
        for block in result.blocks
        if block.block_type == "resource_manifest"
    )
    assert manifest.structured_content["direct_file_count"] == 4
    assert manifest.structured_content["transitive_file_count"] == 5


def test_html_missing_local_resource_is_failed(tmp_path: Path) -> None:
    root = tmp_path / "broken-prototype"
    root.mkdir()
    (root / "index.html").write_text(
        '<html><body><script src="missing.js"></script></body></html>',
        encoding="utf-8",
    )

    result = InputParsingService.default().parse(
        SourceInput(
            source_id="SRC-PROTOTYPE-BROKEN",
            path=root,
            source_kind="prototype",
            authority=Authority.NORMATIVE,
            entry_points=["index.html"],
        )
    )

    assert result.fidelity_report.status == ParseFidelityStatus.FAILED
    assert "input.missing_reference" in {
        item.code for item in result.fidelity_report.errors
    }


def test_image_metadata_is_preserved_but_semantic_reading_remains_partial(
    tmp_path: Path,
) -> None:
    path = tmp_path / "workflow.jpg"
    Image.new("RGB", (17, 9), color="white").save(path)

    result = InputParsingService.default().parse(
        SourceInput(
            source_id="SRC-IMAGE-001",
            path=path,
            source_kind="business_diagram",
            authority=Authority.NORMATIVE,
        )
    )

    assert result.fidelity_report.status == ParseFidelityStatus.PARTIAL
    assert result.fidelity_report.detected_inventory == {
        "images": 1,
        "ocr_blocks": 0,
        "visual_regions": 0,
    }
    image_block = next(
        block for block in result.blocks if block.block_type == "image"
    )
    assert image_block.structured_content["format"] == "JPEG"
    assert image_block.structured_content["width"] == 17
    assert image_block.structured_content["height"] == 9
    assert {item.code for item in result.fidelity_report.unsupported_features} == {
        "input.ocr_unavailable",
        "input.visual_semantics_unavailable",
    }


def test_prototype_source_tree_is_inventoried_without_executing_source(
    tmp_path: Path,
) -> None:
    root = tmp_path / "prototype-source"
    (root / "src").mkdir(parents=True)
    (root / "node_modules").mkdir()
    (root / "package.json").write_text('{"name":"demo"}', encoding="utf-8")
    (root / "src" / "App.tsx").write_text(
        """
function ReviewForm() {
  return (
    <Route path="/reviews">
      <form><button>提交</button></form>
    </Route>
  );
}
""".strip(),
        encoding="utf-8",
    )
    (root / "src" / "api.ts").write_text(
        'export const submit = () => fetch("/api/reviews");',
        encoding="utf-8",
    )
    (root / "node_modules" / "ignored.js").write_text(
        "throw new Error('must not execute');",
        encoding="utf-8",
    )

    result = InputParsingService.default().parse(
        SourceInput(
            source_id="SRC-CODE-001",
            path=root,
            source_kind="prototype_source",
            authority=Authority.TECHNICAL,
        )
    )

    assert result.fidelity_report.status == ParseFidelityStatus.PARTIAL
    assert result.fidelity_report.detected_inventory == {
        "included_files": 3,
        "ignored_files": 1,
        "routes": 1,
        "components": 1,
        "forms": 1,
        "api_calls": 1,
    }
    assert result.fidelity_report.detected_inventory == (
        result.fidelity_report.parsed_inventory
    )
    assert "input.static_source_semantics" in {
        item.code for item in result.fidelity_report.unsupported_features
    }
    assert any(
        block.block_type == "ignored_file"
        and block.locator == "node_modules/ignored.js"
        for block in result.blocks
    )


def test_g0_blocks_required_partial_but_allows_optional_degraded_source(
    tmp_path: Path,
) -> None:
    requirement = tmp_path / "requirement.md"
    requirement.write_text("# 规则\n\n必须提交。", encoding="utf-8")
    diagram = tmp_path / "diagram.png"
    Image.new("RGB", (3, 3), color="white").save(diagram)
    service = InputParsingService.default()
    complete = service.parse(
        SourceInput(
            source_id="SRC-REQ",
            path=requirement,
            source_kind="requirement",
            authority=Authority.NORMATIVE,
            required=True,
        )
    )
    required_partial = service.parse(
        SourceInput(
            source_id="SRC-DIAGRAM-REQUIRED",
            path=diagram,
            source_kind="business_diagram",
            authority=Authority.NORMATIVE,
            required=True,
        )
    )
    optional_partial = service.parse(
        SourceInput(
            source_id="SRC-DIAGRAM-OPTIONAL",
            path=diagram,
            source_kind="business_diagram",
            authority=Authority.HISTORICAL,
            required=False,
        )
    )
    gate = ParseFidelityGate()

    blocked = gate.evaluate([complete, required_partial])
    allowed = gate.evaluate([complete, optional_partial])

    assert blocked.passed is False
    assert blocked.blocked_source_ids == ["SRC-DIAGRAM-REQUIRED"]
    assert "input.required_source_not_complete" in {
        item.code for item in blocked.findings
    }
    assert allowed.passed is True
    assert allowed.blocked_source_ids == []
    assert "input.optional_source_degraded" in {
        item.code for item in allowed.findings
    }


def test_unknown_and_missing_inputs_return_explicit_non_success_status(
    tmp_path: Path,
) -> None:
    unknown = tmp_path / "model.xlsx"
    unknown.write_bytes(b"not an xlsx")
    service = InputParsingService.default()

    unsupported = service.parse(
        SourceInput(
            source_id="SRC-XLSX",
            path=unknown,
            source_kind="spreadsheet",
            authority=Authority.NORMATIVE,
        )
    )
    missing = service.parse(
        SourceInput(
            source_id="SRC-MISSING",
            path=tmp_path / "missing.md",
            source_kind="requirement",
            authority=Authority.NORMATIVE,
        )
    )

    assert unsupported.fidelity_report.status == ParseFidelityStatus.UNSUPPORTED
    assert missing.fidelity_report.status == ParseFidelityStatus.FAILED


def test_content_addressed_capture_keeps_the_parsed_original_reproducible(
    tmp_path: Path,
) -> None:
    original = tmp_path / "requirement.md"
    original.write_text("# 原始规则\n\n必须确认。", encoding="utf-8")
    artifact_root = tmp_path / "artifacts"
    service = InputParsingService.default(artifact_root=artifact_root)

    result = service.parse(
        SourceInput(
            source_id="SRC-CAPTURED",
            path=original,
            source_kind="requirement",
            authority=Authority.NORMATIVE,
        )
    )
    captured_path = Path(result.source.local_path)
    original.write_text("# 被修改的规则", encoding="utf-8")

    assert captured_path.is_file()
    assert artifact_root in captured_path.parents
    assert captured_path.read_text(encoding="utf-8") == "# 原始规则\n\n必须确认。"
    assert result.source.origin_uri == str(original.resolve())


def test_g0_rechecks_inventory_instead_of_trusting_a_complete_label(
    tmp_path: Path,
) -> None:
    path = tmp_path / "requirement.md"
    path.write_text("# 规则\n\n必须提交。", encoding="utf-8")
    parsed = InputParsingService.default().parse(
        SourceInput(
            source_id="SRC-TAMPERED",
            path=path,
            source_kind="requirement",
            authority=Authority.NORMATIVE,
        )
    )
    tampered_report = parsed.fidelity_report.model_copy(
        update={
            "status": ParseFidelityStatus.COMPLETE,
            "parsed_inventory": {
                **parsed.fidelity_report.parsed_inventory,
                "paragraphs": 0,
            },
        }
    )
    tampered = parsed.model_copy(update={"fidelity_report": tampered_report})

    decision = ParseFidelityGate().evaluate([tampered])

    assert decision.passed is False
    assert decision.blocked_source_ids == ["SRC-TAMPERED"]
    assert "input.inventory_mismatch" in {
        item.code for item in decision.findings
    }
